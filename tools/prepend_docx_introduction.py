from pathlib import Path
import sys

from docx import Document


def clone_paragraph(source, target):
    new_paragraph = target.add_paragraph()
    if source.style:
        new_paragraph.style = source.style
    for run in source.runs:
        new_run = new_paragraph.add_run(run.text)
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
    return new_paragraph


def main() -> int:
    if len(sys.argv) != 2:
        print("Usage: prepend_docx_introduction.py <docx-path>", file=sys.stderr)
        return 2

    path = Path(sys.argv[1])
    if not path.exists():
        print(f"File not found: {path}", file=sys.stderr)
        return 1

    existing = Document(path)
    new_doc = Document()

    new_doc.add_heading("Ներածություն", level=1)

    sections = [
        (
            "Թեմայի արդիականությունը",
            [
                (
                    "Ժամանակակից վեբ հավելվածները increasingly կիրառվում են ոչ միայն մեկ կազմակերպության, "
                    "այլ միաժամանակ մի քանի ընկերությունների կամ հաճախորդների սպասարկման համար։ Նման համակարգերը "
                    "հայտնի են որպես բազմաընկերային կամ բազմավարձակալային համակարգեր, որտեղ նույն ծրագրային "
                    "հարթակը պետք է ապահովի տարբեր tenant-ների տվյալների մեկուսացում, ռեսուրսների արդար բաշխում "
                    "և կայուն աշխատանք ծանրաբեռնվածության պայմաններում։"
                ),
                (
                    "Թեմայի արդիականությունը պայմանավորված է նրանով, որ իրական բիզնես համակարգերում հաճախ առաջանում "
                    "են երկարատև հարցումներ՝ հաշվետվությունների գեներացում, մեծածավալ տվյալների ներմուծում, ֆայլերի "
                    "մշակում կամ արտաքին ծառայությունների հետ ինտեգրացիա։ Եթե նման գործողությունները կատարվում են "
                    "սինխրոն ձևով, ապա դրանք կարող են ծանրաբեռնել վեբ սերվերը, մեծացնել response time-ը և բացասաբար "
                    "ազդել մյուս tenant-ների աշխատանքի վրա։ Այդ պատճառով անհրաժեշտ է նախագծել այնպիսի համակարգ, "
                    "որը երկարատև հարցումները կտեղափոխի ասինխրոն մշակման շերտ՝ պահպանելով tenant-ների միջև արդարությունը։"
                ),
            ],
        ),
        (
            "Բազմաընկերային վեբ հավելվածներում երկարատև հարցումների խնդիրը",
            [
                (
                    "Բազմաընկերային միջավայրում երկարատև հարցումների հիմնական խնդիրն այն է, որ տարբեր tenant-ներ "
                    "կիսում են նույն backend ռեսուրսները՝ API server-ներ, queue service, worker microservice-ներ և "
                    "տվյալների պահպանման շերտ։ Եթե մեկ tenant ուղարկում է մեծ քանակությամբ երկարատև job-եր, ապա "
                    "պարզ FIFO հերթի դեպքում այդ tenant-ը կարող է զբաղեցնել ամբողջ worker pool-ը՝ առաջացնելով "
                    "մյուս tenant-ների սպասարկման ուշացում կամ starvation։"
                ),
                (
                    "Այս խնդիրը լուծելու համար անհրաժեշտ է կիրառել tenant-aware queue management մոտեցում, որը "
                    "հաշվի է առնում ոչ միայն job-ի հերթականությունը, այլ նաև այն tenant-ը, որին job-ը պատկանում է։ "
                    "Այդպիսի մոտեցումը պետք է միաժամանակ ապահովի երեք պահանջ՝ երկարատև գործողությունների ասինխրոն "
                    "կատարում, ակտիվ tenant-ների միջև արդար բաշխում և հասանելի worker slot-երի արդյունավետ օգտագործում։"
                ),
            ],
        ),
        (
            "Աշխատանքի նպատակը",
            [
                (
                    "Դիպլոմային աշխատանքի նպատակն է նախագծել և իրականացնել բազմաընկերային վեբ հավելվածների համար "
                    "երկարատև հարցումների ասինխրոն մշակման համակարգ, որը կապահովի job-երի հերթագրում, tenant-ների "
                    "միջև արդար բաշխում, worker slot-երի արդյունավետ օգտագործում, retry/dead-letter մեխանիզմներ և "
                    "տվյալների անվտանգության հիմնական պահանջների ցուցադրական իրականացում։"
                )
            ],
        ),
        (
            "Աշխատանքի խնդիրները",
            [
                "Ուսումնասիրել բազմաընկերային վեբ հավելվածներում երկարատև հարցումների մշակման հիմնախնդիրները։",
                "Նախագծել ասինխրոն job processing համակարգի ընդհանուր ճարտարապետությունը։",
                "Մշակել tenant-aware scheduling ալգորիթմ՝ հիմնված dynamic slot allocation, work-conserving behavior և Round-Robin fairness սկզբունքների վրա։",
                "Իրականացնել Tenant API Layer, Queue Service և Worker Microservice բաղադրիչները .NET տեխնոլոգիաներով։",
                "Ապահովել job payload-ի գաղտնագրում, signed request-ի ցուցադրական մեխանիզմ և tenant isolation-ի հիմք։",
                "Իրականացնել retry, dead-letter և worker failure recovery մեխանիզմներ։",
                "Ստեղծել UI demo dashboard՝ տարբեր tenant ակտիվության և worker slot պայմանների ցուցադրման համար։",
                "Կատարել unit, integration և load testing, ինչպես նաև վերլուծել ստացված արդյունքները։",
            ],
        ),
        (
            "Հետազոտության օբյեկտը և առարկան",
            [
                (
                    "Հետազոտության օբյեկտը բազմաընկերային վեբ հավելվածներում երկարատև հարցումների մշակման գործընթացն է։ "
                    "Հետազոտության առարկան tenant-aware job queue և scheduling մեխանիզմներն են, որոնք ապահովում են "
                    "ասինխրոն մշակման, արդար բաշխման և ռեսուրսների արդյունավետ օգտագործման պահանջները։"
                )
            ],
        ),
        (
            "Օգտագործված տեխնոլոգիաները",
            [
                (
                    "Աշխատանքում օգտագործվել են .NET 8 Worker Service, ASP.NET Core Web API, Entity Framework Core, "
                    "SQL Server, xUnit, Testcontainers և k6 տեխնոլոգիաները։ Ծրագրային իրականացումը ներառում է նաև "
                    "AES հիմքով payload encryption, RSA-PSS հիմքով signed request-ի ստուգում, JSON/file-backed demo storage, "
                    "ինչպես նաև UI dashboard համակարգի աշխատանքի ցուցադրման համար։"
                )
            ],
        ),
        (
            "Աշխատանքի կառուցվածքը",
            [
                (
                    "Աշխատանքը բաղկացած է ներածությունից, չորս գլխից, եզրակացությունից, օգտագործված գրականության ցանկից "
                    "և հավելվածներից։ Առաջին գլխում ներկայացվում են բազմաընկերային վեբ հավելվածներում երկարատև հարցումների "
                    "մշակման հիմնախնդիրները։ Երկրորդ գլխում նկարագրվում է առաջարկվող համակարգի ճարտարապետությունը և "
                    "սեփական scheduling ալգորիթմի նախագծումը։ Երրորդ գլխում ներկայացվում է համակարգի ծրագրային իրականացումը։ "
                    "Չորրորդ գլխում ներկայացվում են փորձարկումները և ստացված արդյունքների վերլուծությունը։"
                )
            ],
        ),
    ]

    for title, paragraphs in sections:
        new_doc.add_heading(title, level=2)
        for text in paragraphs:
            if isinstance(text, str) and len(text) < 180 and not text.endswith("։"):
                new_doc.add_paragraph(text, style="List Bullet")
            else:
                new_doc.add_paragraph(text)

    new_doc.add_page_break()

    for paragraph in existing.paragraphs:
        clone_paragraph(paragraph, new_doc)

    new_doc.save(path)
    print("Prepended introduction successfully.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
